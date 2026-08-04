import io
import json
import zipfile
from collections import Counter
from typing import List
import PyPDF2


def merge_pdfs(pdf_bytes_list: List[bytes]) -> bytes:
    """Merge multiple PDFs into one."""
    writer = PyPDF2.PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def split_pdf(pdf_bytes: bytes, split_at: List[int]) -> List[bytes]:
    """
    Split PDF at given page numbers (1-indexed).
    Returns a list of PDF byte strings.
    split_at=[3, 6] on a 9-page doc → [pages 1-2], [pages 3-5], [pages 6-9]
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    total = len(reader.pages)

    # Build split boundaries (0-indexed start of each chunk)
    boundaries = [0] + [p - 1 for p in sorted(split_at) if 0 < p <= total] + [total]
    boundaries = sorted(set(boundaries))

    result = []
    for i in range(len(boundaries) - 1):
        start = boundaries[i]
        end = boundaries[i + 1]
        writer = PyPDF2.PdfWriter()
        for idx in range(start, end):
            writer.add_page(reader.pages[idx])
        out = io.BytesIO()
        writer.write(out)
        result.append(out.getvalue())
    return result


def _recompress_images(doc, jpeg_quality: int = 55, max_dimension: int = 1600) -> None:
    """
    Re-encode every embedded raster image as JPEG at reduced quality/resolution.
    Embedded images are almost always the dominant contributor to PDF size, so
    this is what actually moves the needle at the "high" compression level —
    stream/object-level cleanup alone barely shrinks image-heavy PDFs.
    """
    from PIL import Image

    seen_xrefs = set()
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                base_image = doc.extract_image(xref)
                pil_img = Image.open(io.BytesIO(base_image["image"]))
                if pil_img.mode not in ("RGB", "L"):
                    pil_img = pil_img.convert("RGB")
                if max(pil_img.size) > max_dimension:
                    ratio = max_dimension / max(pil_img.size)
                    new_size = (max(1, round(pil_img.width * ratio)), max(1, round(pil_img.height * ratio)))
                    pil_img = pil_img.resize(new_size, Image.LANCZOS)
                recompressed = io.BytesIO()
                pil_img.save(recompressed, format="JPEG", quality=jpeg_quality, optimize=True)
                page.replace_image(xref, stream=recompressed.getvalue())
            except Exception:
                continue  # leave images we can't safely recompress (e.g. masks) untouched


def compress_pdf(pdf_bytes: bytes, level: str = "medium") -> bytes:
    """
    Compress a PDF using PyMuPDF, which can meaningfully shrink files by
    garbage-collecting unused/duplicate objects and deflating streams —
    unlike a plain PyPDF2 page copy, which barely changes size because it
    never touches stream compression or image data.

    - "low"    : lossless structural cleanup — merge duplicate objects, drop
                 unreferenced ones, compact the cross-reference table.
    - "medium" : low + deflate (recompress) streams/fonts/images losslessly,
                 strip links.
    - "high"   : medium + re-encode embedded images as JPEG at reduced
                 quality and capped resolution — the biggest lever for
                 image-heavy PDFs (scans, photo-filled documents).
    """
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if level in ("medium", "high"):
        for page in doc:
            try:
                page.remove_links()
            except Exception:
                pass

    if level == "high":
        _recompress_images(doc)

    output = io.BytesIO()
    doc.save(
        output,
        garbage=4,
        deflate=True,
        deflate_images=level in ("medium", "high"),
        deflate_fonts=level in ("medium", "high"),
        clean=level in ("medium", "high"),
    )
    result = output.getvalue()

    # Never hand back a file bigger than the original (can happen on PDFs
    # that are already well-optimized, e.g. re-compressing at "low").
    return result if len(result) < len(pdf_bytes) else pdf_bytes


def extract_pages(pdf_bytes: bytes, pages: List[int]) -> bytes:
    """Extract specific pages (1-indexed) from PDF."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    for p in pages:
        idx = p - 1
        if 0 <= idx < total:
            writer.add_page(reader.pages[idx])
    if len(writer.pages) == 0:
        raise ValueError(f"None of the requested pages exist in this document (it has {total} page{'s' if total != 1 else ''}).")
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def delete_pages(pdf_bytes: bytes, pages: List[int]) -> bytes:
    """Delete specific pages (1-indexed) from PDF."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    pages_to_delete = set(p - 1 for p in pages)
    for idx in range(total):
        if idx not in pages_to_delete:
            writer.add_page(reader.pages[idx])
    if len(writer.pages) == 0:
        raise ValueError("Cannot delete every page — at least one page must remain in the PDF.")
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def rearrange_pages(pdf_bytes: bytes, new_order: List[int]) -> bytes:
    """
    Rearrange pages in the given order (1-indexed).
    new_order=[3,1,2] → page 3 becomes first, page 1 second, page 2 third.
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    for p in new_order:
        idx = p - 1
        if 0 <= idx < total:
            writer.add_page(reader.pages[idx])
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def rotate_pages(pdf_bytes: bytes, pages: List[int], angle: int) -> bytes:
    """
    Rotate specific pages (1-indexed) by angle (90, 180, 270).
    If pages is empty, rotates all pages.
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    pages_to_rotate = set(p - 1 for p in pages) if pages else set(range(total))
    for idx in range(total):
        page = reader.pages[idx]
        if idx in pages_to_rotate:
            page = page.rotate(angle)
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def duplicate_pages(pdf_bytes: bytes, pages: List[int], times: int = 1) -> bytes:
    """
    Duplicate specific pages (1-indexed) inserting copies right after each.
    times=1 means each page appears twice total.
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    total = len(reader.pages)
    pages_to_dup = set(p - 1 for p in pages) if pages else set(range(total))
    for idx in range(total):
        writer.add_page(reader.pages[idx])
        if idx in pages_to_dup:
            for _ in range(times):
                writer.add_page(reader.pages[idx])
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def reverse_page_order(pdf_bytes: bytes) -> bytes:
    """Reverse the page order of a PDF."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    for page in reversed(reader.pages):
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def insert_blank_pages(pdf_bytes: bytes, positions: List[int]) -> bytes:
    """
    Insert blank pages before the given positions (1-indexed).
    position=1 inserts a blank page at the very beginning.
    position=total+1 appends a blank page at the end.
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    total = len(reader.pages)
    if total == 0:
        raise ValueError("Cannot insert blank pages into an empty PDF.")

    # Get dimensions from first page for blank page sizing
    first_page = reader.pages[0]
    width = float(first_page.mediabox.width)
    height = float(first_page.mediabox.height)

    writer = PyPDF2.PdfWriter()

    # Count occurrences per 0-indexed insertion point (not a set) so a
    # repeated position like "2,2" inserts two blanks instead of collapsing
    # to one.
    insert_counts = Counter(p - 1 for p in positions if p <= total)
    # Also handle appending (position = total + 1)
    append_count = sum(1 for p in positions if p > total)

    for idx in range(total):
        for _ in range(insert_counts.get(idx, 0)):
            writer.add_blank_page(width=width, height=height)
        writer.add_page(reader.pages[idx])

    for _ in range(append_count):
        writer.add_blank_page(width=width, height=height)

    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def add_pdf_to_existing(base_pdf_bytes: bytes, new_pdf_bytes: bytes, position: int) -> bytes:
    """
    Insert new_pdf into base_pdf at the given position (1-indexed).
    position=1 → insert at beginning.
    position=total+1 → append at end.
    """
    base_reader = PyPDF2.PdfReader(io.BytesIO(base_pdf_bytes))
    new_reader = PyPDF2.PdfReader(io.BytesIO(new_pdf_bytes))
    writer = PyPDF2.PdfWriter()

    total = len(base_reader.pages)
    insert_at = max(0, min(position - 1, total))

    # Pages before insertion point
    for idx in range(insert_at):
        writer.add_page(base_reader.pages[idx])

    # Inserted PDF pages
    for page in new_reader.pages:
        writer.add_page(page)

    # Remaining base pages
    for idx in range(insert_at, total):
        writer.add_page(base_reader.pages[idx])

    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def get_pdf_info(pdf_bytes: bytes) -> dict:
    """Return basic info about a PDF (page count, etc.)."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    if not reader.pages:
        return {"page_count": 0, "width": 0, "height": 0}
    first_page = reader.pages[0]
    return {
        "page_count": len(reader.pages),
        "width": float(first_page.mediabox.width),
        "height": float(first_page.mediabox.height),
    }


def pdf_to_images(pdf_bytes: bytes, image_format: str = "png") -> List[tuple]:
    """
    Convert all pages of a PDF to images.
    Returns a list of tuples containing (filename, image_bytes).
    """
    import fitz  # PyMuPDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    image_list = []

    for i in range(len(doc)):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes(image_format)
        image_list.append((f"page_{i+1}.{image_format}", img_bytes))

    return image_list


def images_to_pdf(image_bytes_list: List[bytes]) -> bytes:
    """Convert a list of images to a single combined PDF."""
    from PIL import Image
    images = []

    for img_bytes in image_bytes_list:
        img = Image.open(io.BytesIO(img_bytes))
        has_transparency = img.mode in ("RGBA", "LA") or (
            img.mode == "P" and "transparency" in img.info
        )
        if has_transparency:
            # Plain .convert("RGB") drops the alpha channel but keeps
            # whatever RGB values sit underneath it, which are black for
            # freshly-created transparent pixels — so transparent PNG
            # backgrounds (logos, stickers, screenshots) turn solid black
            # instead of staying white. Composite onto white first.
            img = img.convert("RGBA")
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1])
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)

    if not images:
        raise ValueError("At least one image is required.")

    output = io.BytesIO()
    # Save all images merged sequentially into a single PDF stream
    images[0].save(output, format="PDF", save_all=True, append_images=images[1:])
    return output.getvalue()


def _block_remote_resources(uri: str, rel: str):
    """
    xhtml2pdf resolves every resource it finds in the HTML (img src, link
    href, ...) by fetching it, including remote http(s) URLs — an SSRF vector
    when the HTML comes from an untrusted uploaded document (e.g. a Word doc
    with a "linked" rather than embedded image pointing at an internal URL).
    Refusing to resolve anything keeps the conversion sandboxed to layout/text;
    the resource is simply omitted from the output PDF instead of being fetched.
    """
    return None


def word_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert .docx file bytes to a styled PDF using mammoth & xhtml2pdf."""
    import mammoth
    from xhtml2pdf import pisa

    # Convert DOCX layout markup to clean HTML
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    html_content = result.value

    # We xhtml2pdf compile the HTML markup to PDF stream
    pdf_io = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=pdf_io, link_callback=_block_remote_resources)

    if pisa_status.err:
        raise RuntimeError("Failed to convert HTML template layout to PDF.")

    return pdf_io.getvalue()


def unlock_pdf(pdf_bytes: bytes, password: str = "") -> bytes:
    """
    Remove password protection from an encrypted PDF.
    Tries the provided password first, then an empty string as owner password.
    Returns the decrypted PDF bytes.
    Raises ValueError if the PDF is not encrypted or the password is incorrect.
    """
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))

    if not reader.is_encrypted:
        raise ValueError("This PDF is not password-protected.")

    # Try decrypting with the provided password; PyPDF2 tries both user & owner
    result = reader.decrypt(password)
    if result == 0:
        raise ValueError("Incorrect password. Please provide the correct PDF password.")

    # Re-write all pages to a new writer — this strips the encryption entirely
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def pdf_to_word(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF to editable .docx Word file.

    Uses PyMuPDF for structured extraction (text with formatting, images,
    tables) and python-docx for document construction.

    Key capabilities:
    - Preserves text formatting (bold, italic, font size, color)
    - Auto-detects headings from font size and merges wrapped heading lines
    - Converts bullet-point symbols (Symbol font) to real bullet characters
    - Detects chart/graph regions (vector drawings + axis labels) and renders
      them as rasterized image snapshots instead of scattered label text
    - Extracts and embeds real images; skips full-page background images
    - Extracts tables with Table Grid styling
    """
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    word_doc = Document()

    # Set a clean default style
    normal_style = word_doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(4)

    last_table_info = None  # Track table continuations across page boundaries

    for page_idx in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_idx)
        page_rect = page.rect

        if page_idx > 0:
            word_doc.add_page_break()

        # ── Identify table regions so we don't double-emit their text ──
        table_finder = page.find_tables()
        table_rects = []
        table_data_list = []
        for table in table_finder.tables:
            if not _is_false_positive_table(table, page_rect):
                table_rects.append(fitz.Rect(table.bbox))
                table_data_list.append(table.extract())

        # ── Detect chart/graph regions ──
        # Charts are typically: many vector drawings concentrated in an area
        # + small scattered text labels (axis labels, legend entries).
        # We rasterize these regions as images instead of emitting garbled text.
        chart_rects = _detect_chart_regions(page, table_rects)

        # Combine all exclusion zones (tables + charts)
        exclusion_rects = table_rects + chart_rects

        # ── Extract structured text (blocks → lines → spans) ──
        text_dict = page.get_text("dict", sort=True)
        blocks = text_dict.get("blocks", [])

        # Calculate page margin (leftmost text x0 coordinate)
        all_text_x0s = [
            s["bbox"][0]
            for b in blocks
            if b["type"] == 0
            for l in b.get("lines", [])
            for s in l.get("spans", [])
            if s.get("text", "").strip()
        ]
        page_margin = min(all_text_x0s) if all_text_x0s else 56.7

        # Collect all content items sorted by vertical position
        content_items = []

        for block in blocks:
            block_rect = fitz.Rect(block["bbox"])

            # Skip blocks inside table or chart regions
            if any(block_rect.intersects(er) for er in exclusion_rects):
                continue

            y_pos = block["bbox"][1]

            if block["type"] == 0:  # text block
                content_items.append(("text", y_pos, block))
            elif block["type"] == 1:  # image block
                # Skip full-page background/decorative images ONLY IF page has text blocks
                # (preserves scanned/image-only PDFs like large-doc.pdf)
                img_area = block_rect.width * block_rect.height
                page_area = page_rect.width * page_rect.height
                num_text_blocks = sum(1 for b in blocks if b["type"] == 0)
                if img_area > page_area * 0.80 and num_text_blocks > 0:
                    continue
                content_items.append(("image", y_pos, block))

        # Insert tables at their vertical position
        for rect, data in zip(table_rects, table_data_list):
            content_items.append(("table", rect.y0, (rect, data)))

        # Insert chart snapshots at their vertical position
        for chart_rect in chart_rects:
            content_items.append(("chart", chart_rect.y0, (page, chart_rect)))

        # Sort by vertical position for natural reading order
        content_items.sort(key=lambda item: item[1])

        # ── Merge consecutive heading blocks that are really one heading ──
        content_items = _merge_heading_blocks(content_items)

        # ── Render each content item ──
        for item_type, _, payload in content_items:
            if item_type == "text":
                _render_text_block(word_doc, payload, page_margin, page_rect)
            elif item_type == "image":
                _render_image_block(word_doc, payload)
            elif item_type == "table":
                rect, data = payload
                num_cols = max(len(r) for r in data) if data else 0

                is_continuation = False
                if last_table_info is not None and num_cols > 0:
                    if (
                        num_cols == last_table_info["cols"]
                        and abs(rect.x0 - last_table_info["x0"]) < 30
                        and abs(rect.x1 - last_table_info["x1"]) < 30
                    ):
                        is_continuation = True

                if is_continuation:
                    _render_table(
                        word_doc, data, existing_table=last_table_info["table_obj"]
                    )
                else:
                    word_table = _render_table(word_doc, data)
                    if word_table is not None and num_cols > 0:
                        last_table_info = {
                            "cols": num_cols,
                            "x0": rect.x0,
                            "x1": rect.x1,
                            "table_obj": word_table,
                        }
            elif item_type == "chart":
                _render_chart_snapshot(word_doc, payload[0], payload[1])

    output = io.BytesIO()
    word_doc.save(output)
    return output.getvalue()


# ── Helpers for pdf_to_word ──────────────────────────────────────────────────


_BULLET_CHARS = {"\uf0b7", "\uf0a7", "\u2022", "\u25cf", "\u25cb", "\u25aa"}


def _is_bullet_span(span: dict) -> bool:
    """Check if a span is a bullet-point symbol (Symbol/Wingdings font)."""
    font = span.get("font", "").lower()
    text = span.get("text", "").strip()
    if font in ("symbol", "wingdings", "zapfdingbats"):
        return True
    if text in _BULLET_CHARS:
        return True
    return False


def _get_block_heading_level(block: dict) -> int:
    """Return heading level (1-3) for a text block, or 0 for normal text."""
    lines = block.get("lines", [])
    if not lines:
        return 0
    sizes = [
        span["size"]
        for line in lines
        for span in line.get("spans", [])
        if span.get("text", "").strip()
    ]
    if not sizes:
        return 0
    avg = sum(sizes) / len(sizes)
    if avg >= 22:
        return 1
    if avg >= 17:
        return 2
    if avg >= 14:
        return 3
    return 0


def _merge_heading_blocks(content_items: list) -> list:
    """
    Merge consecutive text blocks that are at the same heading level.
    This fixes multi-line headings that PyMuPDF splits into separate blocks.
    """
    if not content_items:
        return content_items

    merged = [content_items[0]]

    for item in content_items[1:]:
        prev = merged[-1]

        # Only merge if both are text blocks
        if prev[0] == "text" and item[0] == "text":
            prev_level = _get_block_heading_level(prev[2])
            curr_level = _get_block_heading_level(item[2])

            # Merge if both are headings at the same level, and they're close
            # vertically (within ~30 pts, typical line spacing for headings)
            if prev_level > 0 and prev_level == curr_level:
                prev_bottom = prev[2]["bbox"][3]
                curr_top = item[2]["bbox"][1]
                if curr_top - prev_bottom < 30:
                    # Merge: append current block's lines to previous block
                    merged_block = dict(prev[2])
                    merged_block["lines"] = list(prev[2].get("lines", [])) + list(item[2].get("lines", []))
                    merged_block["bbox"] = (
                        min(prev[2]["bbox"][0], item[2]["bbox"][0]),
                        prev[2]["bbox"][1],
                        max(prev[2]["bbox"][2], item[2]["bbox"][2]),
                        item[2]["bbox"][3],
                    )
                    merged[-1] = (prev[0], prev[1], merged_block)
                    continue

        merged.append(item)

    return merged


def _detect_chart_regions(page, table_rects: list) -> list:
    """
    Detect chart/graph regions on a page.

    A chart region is identified by a dense cluster of vector drawings
    (bars, lines, axes) combined with small scattered text labels.
    We filter out decorative elements (page backgrounds, thin separator
    lines) to avoid falsely treating entire pages as charts.
    """
    import fitz

    drawings = page.get_drawings()
    if len(drawings) < 10:
        return []

    page_rect = page.rect
    page_area = page_rect.width * page_rect.height

    # ── Filter out non-chart drawings ──
    chart_candidate_rects = []
    for d in drawings:
        r = d.get("rect")
        if not r:
            continue
        rect = fitz.Rect(r)
        rect_area = rect.width * rect.height

        # Skip page-background rectangles (>50% of page)
        if rect_area > page_area * 0.50:
            continue

        # Skip thin lines / hairlines (likely underlines, separators, borders)
        if rect.width < 2 and rect.height < 2:
            continue

        chart_candidate_rects.append(rect)

    # Need a meaningful cluster of drawing primitives
    if len(chart_candidate_rects) < 8:
        return []

    # Find the bounding box of candidate drawings
    union = chart_candidate_rects[0]
    for r in chart_candidate_rects[1:]:
        union = union | r

    # The chart region should be a moderate portion of the page
    chart_area = union.width * union.height
    if chart_area < page_area * 0.03 or chart_area > page_area * 0.70:
        return []

    # Check density: enough drawings concentrated in this region
    contained = sum(1 for r in chart_candidate_rects if union.contains(r))
    if contained < 8:
        return []

    # Skip if this region heavily overlaps with a table
    for tr in table_rects:
        overlap = union & tr
        if overlap.is_empty:
            continue
        overlap_ratio = (overlap.width * overlap.height) / chart_area
        if overlap_ratio > 0.5:
            return []

    # Expand slightly to capture axis labels just outside the chart
    expanded = fitz.Rect(
        union.x0 - 20,
        union.y0 - 20,
        union.x1 + 20,
        union.y1 + 20,
    )
    expanded = expanded & page_rect  # clamp to page bounds

    return [expanded]


def _render_text_block(word_doc, block: dict, page_margin: float = 56.7, page_rect=None) -> None:
    """Add a text block to the Word document with font formatting and precise alignment/indentation."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = block.get("lines", [])
    if not lines:
        return

    # ── Check for bullet points ──
    has_bullet = False
    for line in lines:
        for span in line.get("spans", []):
            if _is_bullet_span(span):
                has_bullet = True
                break
        if has_bullet:
            break

    # Determine heading level
    heading_level = _get_block_heading_level(block)

    if heading_level == 1:
        para = word_doc.add_heading(level=1)
    elif heading_level == 2:
        para = word_doc.add_heading(level=2)
    elif heading_level == 3:
        para = word_doc.add_heading(level=3)
    elif has_bullet:
        para = word_doc.add_paragraph(style="List Bullet")
    else:
        para = word_doc.add_paragraph()

    # ── Auto-center headings if horizontally centered on page ──
    if heading_level > 0 and page_rect is not None:
        block_bbox = block.get("bbox", (0, 0, 0, 0))
        block_center = (block_bbox[0] + block_bbox[2]) / 2.0
        page_center = page_rect.width / 2.0
        if abs(block_center - page_center) < 35:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Calculate exact bullet & paragraph indentation matching PDF ──
    if has_bullet:
        bullet_span = None
        first_text_span = None
        for l in lines:
            for s in l.get("spans", []):
                if not s.get("text", "").strip():
                    continue
                if _is_bullet_span(s):
                    if bullet_span is None:
                        bullet_span = s
                else:
                    if first_text_span is None:
                        first_text_span = s

        if bullet_span and first_text_span:
            b_x0 = bullet_span["bbox"][0]
            t_x0 = first_text_span["bbox"][0]
            left_indent = max(0, t_x0 - page_margin)
            hanging_indent = -(t_x0 - b_x0)
            para.paragraph_format.left_indent = Pt(left_indent)
            para.paragraph_format.first_line_indent = Pt(hanging_indent)
    elif heading_level == 0:
        first_text_span = None
        for l in lines:
            for s in l.get("spans", []):
                if s.get("text", "").strip():
                    first_text_span = s
                    break
            if first_text_span:
                break
        if first_text_span:
            t_x0 = first_text_span["bbox"][0]
            indent = max(0, t_x0 - page_margin)
            if indent > 12:  # Only apply for noticeably indented paragraphs (>12pt)
                para.paragraph_format.left_indent = Pt(indent)

    for line in lines:
        spans = line.get("spans", [])
        for span in spans:
            text = span.get("text", "")

            # Replace bullet symbols with actual bullet character or skip them
            # since we're already using "List Bullet" paragraph style
            if _is_bullet_span(span):
                continue  # skip the bullet symbol; paragraph style handles it

            if not text:
                continue

            run = para.add_run(text)

            # Font size (skip for headings — let Word style control it)
            if heading_level == 0:
                run.font.size = Pt(span.get("size", 11))

            # Bold / italic from flags bitmask
            flags = span.get("flags", 0)
            run.font.bold = bool(flags & 16)    # bit 4
            run.font.italic = bool(flags & 2)   # bit 1

            # Font color (integer RGB packed as 0xRRGGBB)
            color_int = span.get("color", 0)
            if color_int and color_int != 0:
                r = (color_int >> 16) & 0xFF
                g = (color_int >> 8) & 0xFF
                b = color_int & 0xFF
                run.font.color.rgb = RGBColor(r, g, b)

            # Clean up PDF font subset names like "ABCDEF+Arial-Bold" → "Arial"
            font_name = span.get("font", "")
            if font_name:
                clean = font_name.split("+")[-1]  # strip subset prefix
                base = clean.split("-")[0]          # strip style suffix
                if base.lower() not in ("symbol", "wingdings", "zapfdingbats"):
                    run.font.name = base or "Calibri"

        # Add a space between lines within the same block (natural flow)
        if line is not lines[-1]:
            last_text = ""
            for s in reversed(spans):
                last_text = s.get("text", "")
                if last_text:
                    break
            if last_text and not last_text.endswith(" "):
                para.add_run(" ")


def _render_image_block(word_doc, block: dict) -> None:
    """Embed an image block into the Word document."""
    from docx.shared import Inches

    img_data = block.get("image")
    if not img_data:
        return

    try:
        img_stream = io.BytesIO(img_data)
        # Scale to fit within 6-inch page width while keeping aspect ratio
        bbox = block["bbox"]
        img_width_pts = bbox[2] - bbox[0]
        width_inches = min(img_width_pts / 72.0, 6.0)
        width_inches = max(width_inches, 0.5)  # at least half an inch
        word_doc.add_picture(img_stream, width=Inches(width_inches))
    except Exception:
        pass  # skip images we can't decode (e.g. JBIG2 masks)


def _render_chart_snapshot(word_doc, page, chart_rect) -> None:
    """Rasterize a chart/graph region of the page and embed as an image."""
    from docx.shared import Inches

    try:
        # Render just the chart region at 200 DPI for crisp output
        clip = chart_rect
        mat = page.derotation_matrix  # handle rotated pages
        zoom = 200 / 72.0  # 200 DPI
        mat = mat * __import__("fitz").Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, clip=clip)

        img_bytes = pix.tobytes("png")
        img_stream = io.BytesIO(img_bytes)

        # Scale to fit nicely in the document
        width_inches = min(chart_rect.width / 72.0, 6.0)
        word_doc.add_picture(img_stream, width=Inches(width_inches))
    except Exception:
        pass  # if rasterization fails, skip silently


def _is_false_positive_table(table, page_rect) -> bool:
    """
    Check if PyMuPDF's TableFinder detected a false positive table.
    Page borders and decorative outer frames are often misidentified as giant
    single-cell tables containing long narrative paragraphs.
    """
    import fitz
    bbox = fitz.Rect(table.bbox)
    h_ratio = bbox.height / page_rect.height
    area_ratio = (bbox.width * bbox.height) / (page_rect.width * page_rect.height)

    # A table taking >60% of page area and >75% of page height is likely a page border frame
    if area_ratio > 0.60 and h_ratio > 0.75:
        data = table.extract()
        for row in data:
            for cell in row:
                if cell and len(str(cell)) > 150:
                    return True
    return False


def _render_table(word_doc, table_data: list, existing_table=None):
    """Add a table to the Word document, or append rows to an existing table."""
    from docx.shared import Pt

    if not table_data:
        return None

    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data) if table_data else 0
    if num_rows == 0 or num_cols == 0:
        return None

    if existing_table is not None:
        word_table = existing_table
        for row in table_data:
            row_cells = word_table.add_row().cells
            for j, cell_text in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = cell_text or ""
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        return word_table

    word_table = word_doc.add_table(rows=num_rows, cols=num_cols)
    word_table.style = "Table Grid"

    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = word_table.rows[i].cells[j]
                cell.text = cell_text or ""
                # Apply a readable font size to table cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Add a small gap after the table
    word_doc.add_paragraph()
    return word_table


# ─── NEW FEATURES ─────────────────────────────────────────────────────────────


def protect_pdf(pdf_bytes: bytes, password: str) -> bytes:
    """
    Encrypt a PDF with AES-256 password protection using PyMuPDF.
    Both user and owner password are set to the provided password.
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    output = io.BytesIO()
    doc.save(
        output,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw=password,
        owner_pw=password,
        permissions=int(
            fitz.PDF_PERM_ACCESSIBILITY
            | fitz.PDF_PERM_PRINT
            | fitz.PDF_PERM_COPY
            | fitz.PDF_PERM_ANNOTATE
        ),
    )
    return output.getvalue()


def add_watermark(
    pdf_bytes: bytes,
    text: str,
    opacity: float = 0.3,
    angle: float = 45.0,
    font_size: int = 48,
    color: tuple = (0.6, 0.6, 0.6),
) -> bytes:
    """
    Overlay diagonal text watermark on every page of a PDF.
    opacity: 0.0 (invisible) to 1.0 (opaque)
    angle: rotation angle in degrees (any value, not just multiples of 90)
    color: RGB tuple, each 0.0–1.0

    Uses fitz.TextWriter + rotation matrix instead of insert_text(rotate=...)
    because insert_text only accepts 0, 90, 180, 270 for the rotate parameter.
    """
    import fitz
    import math

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page in doc:
        rect = page.rect
        cx = rect.width / 2
        cy = rect.height / 2

        # Build the text writer and measure the text bounding box at origin
        tw_measure = fitz.TextWriter(rect)
        font = fitz.Font("helv")
        tw_measure.append((0, 0), text, font=font, fontsize=font_size)
        text_rect = tw_measure.text_rect
        tw_w = text_rect.width
        tw_h = abs(text_rect.height)

        # Place baseline origin so the text block is visually centered on the page
        # In PyMuPDF, y increases downward; text baseline is at the y coordinate given.
        origin_x = cx - tw_w / 2
        origin_y = cy + tw_h / 4  # slight upward shift to visually center

        # Build the real TextWriter with opacity
        tw = fitz.TextWriter(rect)
        tw.append((origin_x, origin_y), text, font=font, fontsize=font_size)

        # Build a rotation matrix for arbitrary angle around the page center
        rad = math.radians(angle)
        cos_a = math.cos(rad)
        sin_a = math.sin(rad)
        rot_matrix = fitz.Matrix(cos_a, sin_a, -sin_a, cos_a, 0, 0)

        # write_text with morph=(pivot, matrix) rotates text around pivot.
        # color and opacity are valid kwargs on write_text, NOT on append().
        tw.write_text(page, color=color, opacity=opacity, morph=(fitz.Point(cx, cy), rot_matrix))


    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()



def add_page_numbers(
    pdf_bytes: bytes,
    position: str = "bottom-center",
    font_size: int = 10,
    start_number: int = 1,
    prefix: str = "Page",
) -> bytes:
    """
    Stamp page numbers on every page of the PDF.
    position: 'bottom-center' | 'bottom-left' | 'bottom-right' | 'top-center' | 'top-left' | 'top-right'
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    margin = 28

    for i, page in enumerate(doc):
        label = f"{prefix} {i + start_number}" if prefix else str(i + start_number)
        rect = page.rect
        text_width = len(label) * font_size * 0.55  # approximate

        if "bottom" in position:
            y = rect.height - margin
        else:
            y = margin + font_size

        if "center" in position:
            x = (rect.width - text_width) / 2
        elif "left" in position:
            x = margin
        else:
            x = rect.width - text_width - margin

        page.insert_text(
            fitz.Point(x, y),
            label,
            fontsize=font_size,
            color=(0.2, 0.2, 0.2),
            overlay=True,
        )

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def extract_text(pdf_bytes: bytes) -> str:
    """Extract all text from a PDF and return as a plain string."""
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text_parts = []
    for i, page in enumerate(doc):
        text_parts.append(f"--- Page {i + 1} ---\n")
        text_parts.append(page.get_text("text"))
        text_parts.append("\n")
    return "\n".join(text_parts)


def extract_images_from_pdf(pdf_bytes: bytes) -> List[tuple]:
    """
    Extract all embedded images from a PDF.
    Returns a list of (filename, image_bytes) tuples for ZIP packaging.
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    img_count = 0

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)

        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                img_ext = base_image["ext"]
                img_count += 1
                filename = f"page{page_num + 1}_img{img_index + 1}.{img_ext}"
                images.append((filename, img_bytes))
            except Exception:
                continue

    if not images:
        raise ValueError("No images found in this PDF.")

    return images


def pdf_to_excel(pdf_bytes: bytes) -> bytes:
    """
    Extract tables from a PDF and write to an .xlsx file.
    Uses pdfplumber for table extraction and openpyxl for Excel writing.
    """
    import pdfplumber
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)  # Remove default empty sheet
    sheet_count = 0

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if not tables:
                continue

            for table_idx, table in enumerate(tables):
                sheet_count += 1
                ws = wb.create_sheet(title=f"P{page_num + 1}_T{table_idx + 1}")
                for row in table:
                    ws.append([cell if cell is not None else "" for cell in row])

    if sheet_count == 0:
        # If no tables found, extract text into a single sheet
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            ws = wb.create_sheet(title="Extracted Text")
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                ws.append([f"--- Page {page_num + 1} ---"])
                for line in text.split("\n"):
                    ws.append([line])
                ws.append([""])

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def add_signature(
    pdf_bytes: bytes,
    sig_image_bytes: bytes,
    page_num: int = 1,
    x: float = 100.0,
    y: float = 100.0,
    width: float = 200.0,
    height: float = 80.0,
) -> bytes:
    """
    Embed a signature image onto a specific page of the PDF.
    page_num: 1-indexed page number
    x, y: top-left corner position in PDF points
    width, height: size of signature in PDF points
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = len(doc)
    page_idx = max(0, min(page_num - 1, total - 1))
    page = doc.load_page(page_idx)

    rect = fitz.Rect(x, y, x + width, y + height)
    page.insert_image(rect, stream=sig_image_bytes)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def annotate_pdf(pdf_bytes: bytes, annotations: list) -> bytes:
    """
    Add annotations to a PDF.
    Each annotation is a dict with:
      - type: 'highlight' | 'text'
      - page: 1-indexed page number
      - x, y: position in PDF points
      - For 'highlight': x2, y2 (end of highlight rect), color (hex string)
      - For 'text': content (string), width, height
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    def hex_to_rgb(hex_color: str) -> tuple:
        """Convert '#rrggbb' or 'rrggbb' to (r, g, b) floats 0..1."""
        hex_color = hex_color.lstrip("#")
        if len(hex_color) == 3:
            hex_color = "".join(c * 2 for c in hex_color)
        r = int(hex_color[0:2], 16) / 255
        g = int(hex_color[2:4], 16) / 255
        b = int(hex_color[4:6], 16) / 255
        return (r, g, b)

    total = len(doc)

    for ann in annotations:
        page_num = int(ann.get("page", 1))
        page_idx = max(0, min(page_num - 1, total - 1))
        page = doc.load_page(page_idx)
        ann_type = ann.get("type", "text")
        x = float(ann.get("x", 50))
        y = float(ann.get("y", 50))

        if ann_type == "highlight":
            x2 = float(ann.get("x2", x + 200))
            y2 = float(ann.get("y2", y + 20))
            color_hex = ann.get("color", "#FFFF00")
            rgb = hex_to_rgb(color_hex)
            rect = fitz.Rect(x, y, x2, y2)
            # Use rect_annot with opacity for a reliable highlight overlay box
            rect_ann = page.add_rect_annot(rect)
            rect_ann.set_colors(stroke=rgb, fill=rgb)
            rect_ann.set_opacity(0.35)
            rect_ann.update()

        elif ann_type == "text":
            content = ann.get("content", "")
            ann_width = float(ann.get("width", 200))
            ann_height = float(ann.get("height", 60))
            color_hex = ann.get("color", "#FFD700")
            rgb = hex_to_rgb(color_hex)
            rect = fitz.Rect(x, y, x + ann_width, y + ann_height)
            # border_color is omitted to avoid PyMuPDF rich_text ValueError
            text_ann = page.add_freetext_annot(
                rect,
                content,
                fontsize=12,
                text_color=(0, 0, 0),
                fill_color=rgb,
            )
            text_ann.update()

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def redact_pdf(pdf_bytes: bytes, rects: list) -> bytes:
    """
    Permanently remove content (text + graphics) inside the given rectangles
    and draw a solid black box over each — unlike a watermark/annotation
    overlay, the underlying content is deleted, not just visually covered.
    Each rect: {page: 1-indexed page number, x, y, x2, y2} in PDF points.
    """
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = len(doc)
    touched_pages = set()

    for r in rects:
        page_num = int(r.get("page", 1))
        page_idx = max(0, min(page_num - 1, total - 1))
        page = doc.load_page(page_idx)
        x = float(r.get("x", 0))
        y = float(r.get("y", 0))
        x2 = float(r.get("x2", x + 100))
        y2 = float(r.get("y2", y + 20))
        page.add_redact_annot(fitz.Rect(x, y, x2, y2), fill=(0, 0, 0))
        touched_pages.add(page_idx)

    for page_idx in touched_pages:
        # graphics=2 also strips graphics that only partially overlap the
        # rect (the default, graphics=1, only removes fully-contained ones
        # and leaves a partially-covered vector shape intact underneath the
        # black box — defeats the purpose of "not just an overlay").
        doc.load_page(page_idx).apply_redactions(images=2, graphics=2, text=0)

    output = io.BytesIO()
    doc.save(output, garbage=4, deflate=True)
    return output.getvalue()
